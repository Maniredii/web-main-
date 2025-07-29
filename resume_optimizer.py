"""
Resume Optimizer for ATS Compatibility
Intelligently optimizes resumes based on job requirements using AI
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from resume_parser import ResumeContent, ResumeParser
from job_description_parser import JobRequirements, JobDescriptionParser

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class OptimizationResult:
    """Result of resume optimization process"""
    original_resume: ResumeContent
    job_requirements: JobRequirements
    optimized_resume: ResumeContent
    optimization_score: float
    keyword_matches: Dict[str, int]
    improvements_made: List[str]
    optimization_date: str
    output_file_path: str

class ResumeOptimizer:
    """AI-powered resume optimizer for ATS compatibility"""
    
    def __init__(self, ollama_client=None):
        self.ollama_client = ollama_client
        self.resume_parser = ResumeParser(ollama_client)
        self.job_parser = JobDescriptionParser(ollama_client)
        
        # ATS optimization weights
        self.optimization_weights = {
            'keyword_matching': 0.4,
            'skill_alignment': 0.3,
            'experience_relevance': 0.2,
            'formatting_quality': 0.1
        }
        
        # Keyword integration strategies
        self.integration_strategies = [
            'natural_insertion',
            'skill_section_enhancement',
            'summary_optimization',
            'experience_tailoring'
        ]
    
    def optimize_resume_for_job(self, resume_path: str, job_description: str, 
                               job_title: str = "", company: str = "",
                               output_path: str = None) -> Optional[OptimizationResult]:
        """
        Optimize resume for specific job posting
        
        Args:
            resume_path: Path to original resume file
            job_description: Job description text
            job_title: Job title
            company: Company name
            output_path: Path for optimized resume output
            
        Returns:
            OptimizationResult with optimization details
        """
        logger.info(f"Starting resume optimization for: {job_title} at {company}")
        
        try:
            # Parse original resume
            original_resume = self.resume_parser.parse_resume(resume_path)
            if not original_resume:
                logger.error("Failed to parse original resume")
                return None
            
            # Parse job requirements
            job_requirements = self.job_parser.parse_job_description(
                job_description, job_title=job_title, company=company
            )
            
            # Create optimized version
            optimized_resume = self._create_optimized_resume(original_resume, job_requirements)
            
            # Calculate optimization metrics
            optimization_score = self._calculate_optimization_score(optimized_resume, job_requirements)
            keyword_matches = self._analyze_keyword_matches(optimized_resume, job_requirements)
            improvements_made = self._track_improvements(original_resume, optimized_resume)
            
            # Generate output file
            if not output_path:
                base_name = os.path.splitext(os.path.basename(resume_path))[0]
                company_safe = re.sub(r'[^\w\-_]', '_', company) if company else "company"
                output_path = f"{base_name}_optimized_{company_safe}.docx"
            
            self._generate_optimized_document(optimized_resume, output_path)
            
            # Create result object
            result = OptimizationResult(
                original_resume=original_resume,
                job_requirements=job_requirements,
                optimized_resume=optimized_resume,
                optimization_score=optimization_score,
                keyword_matches=keyword_matches,
                improvements_made=improvements_made,
                optimization_date=datetime.now().isoformat(),
                output_file_path=output_path
            )
            
            logger.info(f"Resume optimization completed. Score: {optimization_score:.2f}")
            return result
            
        except Exception as e:
            logger.error(f"Resume optimization failed: {e}")
            return None
    
    def _create_optimized_resume(self, original: ResumeContent, 
                               job_req: JobRequirements) -> ResumeContent:
        """Create optimized version of resume based on job requirements"""
        optimized = ResumeContent()
        
        # Copy basic information
        optimized.name = original.name
        optimized.email = original.email
        optimized.phone = original.phone
        optimized.location = original.location
        optimized.linkedin_url = original.linkedin_url
        optimized.github_url = original.github_url
        optimized.portfolio_url = original.portfolio_url
        
        # Optimize professional summary
        optimized.summary = self._optimize_summary(original.summary, job_req)
        
        # Optimize skills sections
        optimized.technical_skills = self._optimize_technical_skills(original, job_req)
        optimized.programming_languages = self._optimize_programming_languages(original, job_req)
        optimized.frameworks_tools = self._optimize_frameworks_tools(original, job_req)
        optimized.soft_skills = self._optimize_soft_skills(original, job_req)
        
        # Optimize work experience
        optimized.work_experience = self._optimize_work_experience(original.work_experience, job_req)
        
        # Copy other sections with potential enhancements
        optimized.education = original.education
        optimized.certifications = self._enhance_certifications(original.certifications, job_req)
        optimized.projects = self._optimize_projects(original.projects, job_req)
        optimized.achievements = original.achievements
        
        # Set metadata
        optimized.file_path = original.file_path
        optimized.file_type = original.file_type
        optimized.parsed_date = datetime.now().isoformat()
        
        return optimized
    
    def _optimize_summary(self, original_summary: str, job_req: JobRequirements) -> str:
        """Optimize professional summary with job-specific keywords"""
        if not self.ollama_client:
            return self._basic_summary_optimization(original_summary, job_req)
        
        try:
            prompt = f"""
            Optimize this professional summary for the following job:
            
            Job Title: {job_req.title}
            Company: {job_req.company}
            Key Requirements: {', '.join(job_req.required_skills[:10])}
            
            Original Summary:
            {original_summary}
            
            Create an optimized summary that:
            1. Incorporates relevant keywords naturally
            2. Highlights matching skills and experience
            3. Maintains professional tone
            4. Stays under 150 words
            5. Focuses on value proposition for this specific role
            
            Return only the optimized summary text.
            """
            
            response = self.ollama_client.query(prompt, max_tokens=200)
            if response and len(response.strip()) > 50:
                return response.strip()
            
        except Exception as e:
            logger.warning(f"AI summary optimization failed: {e}")
        
        return self._basic_summary_optimization(original_summary, job_req)
    
    def _basic_summary_optimization(self, original_summary: str, job_req: JobRequirements) -> str:
        """Basic summary optimization without AI"""
        if not original_summary:
            # Create basic summary from job requirements
            key_skills = job_req.required_skills[:5]
            return f"Experienced professional with expertise in {', '.join(key_skills)}. " \
                   f"Proven track record in {job_req.experience_level} level positions."
        
        # Add key job requirements to existing summary
        summary = original_summary
        
        # Insert top 3 required skills if not already present
        for skill in job_req.required_skills[:3]:
            if skill.lower() not in summary.lower():
                summary = f"{summary} Experienced with {skill}."
        
        return summary
    
    def _optimize_technical_skills(self, original: ResumeContent, 
                                 job_req: JobRequirements) -> List[str]:
        """Optimize technical skills section"""
        optimized_skills = set(original.technical_skills)
        
        # Add matching required skills
        for skill in job_req.required_skills:
            if self._is_technical_skill(skill):
                optimized_skills.add(skill)
        
        # Add matching preferred skills
        for skill in job_req.preferred_skills:
            if self._is_technical_skill(skill):
                optimized_skills.add(skill)
        
        # Prioritize skills that appear in job requirements
        prioritized_skills = []
        job_keywords = set(job_req.ats_keywords)
        
        # First add skills that match job requirements
        for skill in optimized_skills:
            if skill.lower() in [kw.lower() for kw in job_keywords]:
                prioritized_skills.append(skill)
        
        # Then add remaining skills
        for skill in optimized_skills:
            if skill not in prioritized_skills:
                prioritized_skills.append(skill)
        
        return prioritized_skills
    
    def _optimize_programming_languages(self, original: ResumeContent, 
                                      job_req: JobRequirements) -> List[str]:
        """Optimize programming languages list"""
        optimized_langs = set(original.programming_languages)
        
        # Add languages mentioned in job requirements
        for lang in job_req.programming_languages:
            optimized_langs.add(lang)
        
        # Prioritize based on job requirements
        job_langs = [lang.lower() for lang in job_req.programming_languages]
        prioritized = []
        
        for lang in optimized_langs:
            if lang.lower() in job_langs:
                prioritized.append(lang)
        
        for lang in optimized_langs:
            if lang not in prioritized:
                prioritized.append(lang)
        
        return prioritized
    
    def _optimize_frameworks_tools(self, original: ResumeContent, 
                                 job_req: JobRequirements) -> List[str]:
        """Optimize frameworks and tools list"""
        optimized_tools = set(original.frameworks_tools)
        
        # Add tools mentioned in job requirements
        for tool in job_req.frameworks_tools:
            optimized_tools.add(tool)
        
        return list(optimized_tools)
    
    def _optimize_soft_skills(self, original: ResumeContent, 
                            job_req: JobRequirements) -> List[str]:
        """Optimize soft skills based on job requirements"""
        optimized_soft = set(original.soft_skills)
        
        # Add soft skills from job requirements
        for skill in job_req.soft_skills:
            optimized_soft.add(skill)
        
        return list(optimized_soft)
    
    def _optimize_work_experience(self, original_experience: List[Dict[str, Any]], 
                                job_req: JobRequirements) -> List[Dict[str, Any]]:
        """Optimize work experience descriptions with job-relevant keywords"""
        if not original_experience:
            return []
        
        optimized_experience = []
        
        for job in original_experience:
            optimized_job = job.copy()
            
            # Enhance job description with relevant keywords
            if 'description' in optimized_job:
                optimized_job['description'] = self._enhance_job_description(
                    optimized_job['description'], job_req
                )
            
            optimized_experience.append(optimized_job)
        
        return optimized_experience
    
    def _enhance_job_description(self, description: str, job_req: JobRequirements) -> str:
        """Enhance job description with relevant keywords"""
        if not description:
            return description
        
        enhanced_desc = description
        
        # Add relevant keywords naturally
        relevant_keywords = []
        
        # Find keywords that could naturally fit
        for keyword in job_req.ats_keywords[:10]:  # Top 10 keywords
            if keyword.lower() not in description.lower() and len(keyword) > 2:
                relevant_keywords.append(keyword)
        
        # Add keywords naturally to description
        if relevant_keywords and self.ollama_client:
            try:
                prompt = f"""
                Enhance this job description by naturally incorporating these relevant keywords: {', '.join(relevant_keywords[:5])}
                
                Original description:
                {description}
                
                Return an enhanced version that:
                1. Naturally incorporates 2-3 of the keywords
                2. Maintains the original meaning and tone
                3. Doesn't sound forced or artificial
                4. Keeps the same length approximately
                
                Return only the enhanced description.
                """
                
                response = self.ollama_client.query(prompt, max_tokens=300)
                if response and len(response.strip()) > len(description) * 0.8:
                    return response.strip()
                    
            except Exception as e:
                logger.warning(f"AI description enhancement failed: {e}")
        
        # Basic keyword insertion as fallback
        if relevant_keywords:
            enhanced_desc += f" Utilized {relevant_keywords[0]} in daily operations."
        
        return enhanced_desc
    
    def _enhance_certifications(self, original_certs: List[str], 
                              job_req: JobRequirements) -> List[str]:
        """Enhance certifications list with relevant ones from job requirements"""
        enhanced_certs = original_certs.copy()
        
        # Add certifications mentioned in job requirements
        for cert in job_req.certifications:
            if cert not in enhanced_certs:
                enhanced_certs.append(cert)
        
        return enhanced_certs
    
    def _optimize_projects(self, original_projects: List[Dict[str, str]], 
                         job_req: JobRequirements) -> List[Dict[str, str]]:
        """Optimize project descriptions with job-relevant keywords"""
        if not original_projects:
            return []
        
        optimized_projects = []
        
        for project in original_projects:
            optimized_project = project.copy()
            
            if 'description' in optimized_project:
                optimized_project['description'] = self._enhance_project_description(
                    optimized_project['description'], job_req
                )
            
            optimized_projects.append(optimized_project)
        
        return optimized_projects
    
    def _enhance_project_description(self, description: str, job_req: JobRequirements) -> str:
        """Enhance project description with relevant technologies"""
        if not description:
            return description
        
        # Add relevant technologies mentioned in job requirements
        relevant_techs = []
        for tech in job_req.frameworks_tools + job_req.programming_languages:
            if tech.lower() not in description.lower():
                relevant_techs.append(tech)
        
        if relevant_techs:
            description += f" Technologies used include {relevant_techs[0]}."
        
        return description
    
    def _is_technical_skill(self, skill: str) -> bool:
        """Determine if a skill is technical in nature"""
        technical_indicators = [
            'programming', 'development', 'software', 'framework', 'database',
            'cloud', 'api', 'web', 'mobile', 'data', 'machine learning',
            'ai', 'devops', 'testing', 'security'
        ]
        
        skill_lower = skill.lower()
        return any(indicator in skill_lower for indicator in technical_indicators)
    
    def _calculate_optimization_score(self, optimized_resume: ResumeContent, 
                                    job_req: JobRequirements) -> float:
        """Calculate optimization score based on keyword matching and relevance"""
        score = 0.0
        
        # Keyword matching score
        keyword_score = self._calculate_keyword_score(optimized_resume, job_req)
        score += keyword_score * self.optimization_weights['keyword_matching']
        
        # Skill alignment score
        skill_score = self._calculate_skill_alignment_score(optimized_resume, job_req)
        score += skill_score * self.optimization_weights['skill_alignment']
        
        # Experience relevance score
        exp_score = self._calculate_experience_relevance_score(optimized_resume, job_req)
        score += exp_score * self.optimization_weights['experience_relevance']
        
        # Formatting quality score
        format_score = optimized_resume.ats_score if optimized_resume.ats_score else 70.0
        score += format_score * self.optimization_weights['formatting_quality']
        
        return min(score, 100.0)
    
    def _calculate_keyword_score(self, resume: ResumeContent, job_req: JobRequirements) -> float:
        """Calculate keyword matching score"""
        if not job_req.ats_keywords:
            return 50.0
        
        resume_text = self._get_resume_text(resume).lower()
        matched_keywords = 0
        
        for keyword in job_req.ats_keywords:
            if keyword.lower() in resume_text:
                matched_keywords += 1
        
        return (matched_keywords / len(job_req.ats_keywords)) * 100
    
    def _calculate_skill_alignment_score(self, resume: ResumeContent, 
                                       job_req: JobRequirements) -> float:
        """Calculate skill alignment score"""
        resume_skills = set()
        resume_skills.update([s.lower() for s in resume.technical_skills])
        resume_skills.update([s.lower() for s in resume.programming_languages])
        resume_skills.update([s.lower() for s in resume.frameworks_tools])
        
        job_skills = set()
        job_skills.update([s.lower() for s in job_req.required_skills])
        job_skills.update([s.lower() for s in job_req.programming_languages])
        job_skills.update([s.lower() for s in job_req.frameworks_tools])
        
        if not job_skills:
            return 50.0
        
        matched_skills = len(resume_skills.intersection(job_skills))
        return (matched_skills / len(job_skills)) * 100
    
    def _calculate_experience_relevance_score(self, resume: ResumeContent, 
                                            job_req: JobRequirements) -> float:
        """Calculate experience relevance score"""
        if not resume.work_experience:
            return 30.0
        
        # Basic scoring based on presence of experience
        score = 60.0
        
        # Bonus for relevant keywords in experience descriptions
        exp_text = ""
        for job in resume.work_experience:
            if 'description' in job:
                exp_text += job['description'].lower() + " "
        
        keyword_matches = 0
        for keyword in job_req.ats_keywords[:10]:  # Check top 10 keywords
            if keyword.lower() in exp_text:
                keyword_matches += 1
        
        score += (keyword_matches / 10) * 40  # Up to 40 bonus points
        
        return min(score, 100.0)
    
    def _analyze_keyword_matches(self, resume: ResumeContent, 
                               job_req: JobRequirements) -> Dict[str, int]:
        """Analyze keyword matches between resume and job requirements"""
        resume_text = self._get_resume_text(resume).lower()
        keyword_matches = {}
        
        for keyword in job_req.ats_keywords:
            count = resume_text.count(keyword.lower())
            if count > 0:
                keyword_matches[keyword] = count
        
        return keyword_matches
    
    def _track_improvements(self, original: ResumeContent, 
                          optimized: ResumeContent) -> List[str]:
        """Track improvements made during optimization"""
        improvements = []
        
        # Check for added skills
        original_skills = set(original.technical_skills + original.programming_languages + 
                            original.frameworks_tools)
        optimized_skills = set(optimized.technical_skills + optimized.programming_languages + 
                             optimized.frameworks_tools)
        
        new_skills = optimized_skills - original_skills
        if new_skills:
            improvements.append(f"Added {len(new_skills)} relevant technical skills")
        
        # Check for summary optimization
        if len(optimized.summary) > len(original.summary):
            improvements.append("Enhanced professional summary with job-specific keywords")
        
        # Check for experience enhancement
        if optimized.work_experience and original.work_experience:
            if len(str(optimized.work_experience)) > len(str(original.work_experience)):
                improvements.append("Enhanced work experience descriptions")
        
        return improvements
    
    def _get_resume_text(self, resume: ResumeContent) -> str:
        """Get all text content from resume for analysis"""
        text_parts = [
            resume.name, resume.summary, resume.objective,
            ' '.join(resume.technical_skills),
            ' '.join(resume.programming_languages),
            ' '.join(resume.frameworks_tools),
            ' '.join(resume.soft_skills)
        ]
        
        # Add work experience text
        for job in resume.work_experience:
            if isinstance(job, dict):
                text_parts.extend([
                    job.get('title', ''),
                    job.get('company', ''),
                    job.get('description', '')
                ])
        
        return ' '.join([part for part in text_parts if part])
    
    def _generate_optimized_document(self, resume: ResumeContent, output_path: str):
        """Generate optimized resume document"""
        try:
            doc = docx.Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add header with contact information
            header = doc.add_heading(resume.name, 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact_info = []
            if resume.email:
                contact_info.append(resume.email)
            if resume.phone:
                contact_info.append(resume.phone)
            if resume.location:
                contact_info.append(resume.location)
            
            if contact_info:
                contact_para = doc.add_paragraph(' | '.join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # URLs
            url_info = []
            if resume.linkedin_url:
                url_info.append(resume.linkedin_url)
            if resume.github_url:
                url_info.append(resume.github_url)
            
            if url_info:
                url_para = doc.add_paragraph(' | '.join(url_info))
                url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Professional Summary
            if resume.summary:
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(resume.summary)
            
            # Technical Skills
            if resume.technical_skills or resume.programming_languages or resume.frameworks_tools:
                doc.add_heading('Technical Skills', level=1)
                
                if resume.programming_languages:
                    doc.add_paragraph(f"Programming Languages: {', '.join(resume.programming_languages)}")
                
                if resume.frameworks_tools:
                    doc.add_paragraph(f"Frameworks & Tools: {', '.join(resume.frameworks_tools)}")
                
                if resume.technical_skills:
                    doc.add_paragraph(f"Technical Skills: {', '.join(resume.technical_skills)}")
            
            # Work Experience
            if resume.work_experience:
                doc.add_heading('Professional Experience', level=1)
                
                for job in resume.work_experience:
                    if isinstance(job, dict):
                        # Job title and company
                        title_company = f"{job.get('title', '')} - {job.get('company', '')}"
                        job_heading = doc.add_paragraph(title_company)
                        job_heading.style = 'Heading 2'
                        
                        # Duration
                        if job.get('duration'):
                            doc.add_paragraph(job['duration'])
                        
                        # Description
                        if job.get('description'):
                            doc.add_paragraph(job['description'])
                        
                        doc.add_paragraph()  # Add space
            
            # Education
            if resume.education:
                doc.add_heading('Education', level=1)
                for edu in resume.education:
                    if isinstance(edu, dict):
                        edu_text = f"{edu.get('degree', '')} in {edu.get('field', '')}"
                        if edu.get('institution'):
                            edu_text += f" - {edu['institution']}"
                        if edu.get('year'):
                            edu_text += f" ({edu['year']})"
                        doc.add_paragraph(edu_text)
            
            # Projects
            if resume.projects:
                doc.add_heading('Projects', level=1)
                for project in resume.projects:
                    if isinstance(project, dict):
                        project_name = project.get('name', '')
                        if project_name:
                            proj_heading = doc.add_paragraph(project_name)
                            proj_heading.style = 'Heading 3'
                        
                        if project.get('description'):
                            doc.add_paragraph(project['description'])
            
            # Certifications
            if resume.certifications:
                doc.add_heading('Certifications', level=1)
                for cert in resume.certifications:
                    doc.add_paragraph(f"• {cert}")
            
            # Save document
            doc.save(output_path)
            logger.info(f"Optimized resume saved to: {output_path}")
            
        except Exception as e:
            logger.error(f"Failed to generate optimized document: {e}")

# Example usage
if __name__ == "__main__":
    optimizer = ResumeOptimizer()
    
    # Test optimization
    sample_job_desc = """
    We are seeking a Python Developer with 3+ years of experience.
    Required skills: Python, Django, PostgreSQL, AWS, Git
    Preferred: React, Docker, Agile methodologies
    """
    
    result = optimizer.optimize_resume_for_job(
        "sample resume.docx",
        sample_job_desc,
        "Python Developer",
        "Tech Company"
    )
    
    if result:
        print(f"Optimization Score: {result.optimization_score:.2f}")
        print(f"Keyword Matches: {len(result.keyword_matches)}")
        print(f"Improvements: {result.improvements_made}")
