#!/usr/bin/env python3
"""
LinkedIn Job Applier GUI with Resume Upload
Complete automation interface for applying to jobs on LinkedIn
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import threading
import json
import os
import time
from datetime import datetime
from simple_puppeteer_bridge import PuppeteerBridge
from docx import Document
import shutil

class LinkedInJobApplierGUI:
    """Main GUI for LinkedIn job automation with resume management"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("🚀 LinkedIn Job Applier Pro")
        self.root.geometry("900x700")
        self.root.configure(bg='#f0f0f0')
        
        # Initialize components
        self.puppeteer_bridge = PuppeteerBridge()
        self.resume_path = None
        self.resume_text = ""
        self.is_automation_running = False
        
        # Setup UI
        self.setup_ui()
        self.load_saved_settings()
        
    def setup_ui(self):
        """Setup the complete user interface"""
        # Scrollable container (so sections below Start button are visible)
        container = ttk.Frame(self.root)
        container.grid(row=0, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))

        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        canvas = tk.Canvas(container, highlightthickness=0)
        vscroll = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vscroll.set)

        scrollable_frame = ttk.Frame(canvas, padding="20")

        # When the size of the frame changes, update the scrollregion of the canvas
        scrollable_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        # Create window inside canvas
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        # Ensure canvas expands
        container.columnconfigure(0, weight=1)
        container.rowconfigure(0, weight=1)
        canvas.grid(row=0, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))
        vscroll.grid(row=0, column=1, sticky=(tk.N, tk.S))

        # Mousewheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        # Make inner frame width track canvas width
        def _resize_frame(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind("<Configure>", _resize_frame)

        # Build UI into the scrollable frame
        self.create_header(scrollable_frame)
        self.create_resume_section(scrollable_frame)
        self.create_job_search_section(scrollable_frame)
        self.create_automation_controls(scrollable_frame)
        self.create_status_section(scrollable_frame)
        self.create_results_section(scrollable_frame)
        self.create_settings_section(scrollable_frame)
        
    def create_header(self, parent):
        """Create the header section"""
        header_frame = ttk.Frame(parent)
        header_frame.grid(row=0, column=0, columnspan=2, pady=(0, 20), sticky=(tk.W, tk.E))
        
        # Title
        title_label = ttk.Label(header_frame, 
                               text="🚀 LinkedIn Job Applier Pro", 
                               font=("Arial", 20, "bold"))
        title_label.pack()
        
        # Subtitle
        subtitle_label = ttk.Label(header_frame, 
                                  text="Automated job searching and application with AI-powered resume optimization", 
                                  font=("Arial", 12))
        subtitle_label.pack(pady=(5, 0))
        
    def create_resume_section(self, parent):
        """Create the resume upload and management section"""
        resume_frame = ttk.LabelFrame(parent, text="📄 Resume Management", padding="15")
        resume_frame.grid(row=1, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Resume file selection
        ttk.Label(resume_frame, text="Resume File (.docx):").grid(row=0, column=0, sticky=tk.W, pady=5)
        
        self.resume_path_var = tk.StringVar()
        resume_entry = ttk.Entry(resume_frame, textvariable=self.resume_path_var, width=50, state='readonly')
        resume_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(10, 10), pady=5)
        
        browse_btn = ttk.Button(resume_frame, text="Browse", command=self.browse_resume)
        browse_btn.grid(row=0, column=2, pady=5)
        
        # Resume preview
        ttk.Label(resume_frame, text="Resume Preview:").grid(row=1, column=0, sticky=tk.W, pady=(15, 5))
        
        self.resume_preview = scrolledtext.ScrolledText(resume_frame, height=6, width=70, wrap=tk.WORD)
        self.resume_preview.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # Resume actions
        resume_actions_frame = ttk.Frame(resume_frame)
        resume_actions_frame.grid(row=3, column=0, columnspan=3, pady=(10, 0))
        
        ttk.Button(resume_actions_frame, text="📝 Edit Resume", command=self.edit_resume).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(resume_actions_frame, text="💾 Save Resume", command=self.save_resume).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(resume_actions_frame, text="🔄 Refresh Preview", command=self.refresh_resume_preview).pack(side=tk.LEFT)
        
        resume_frame.columnconfigure(1, weight=1)
        
    def create_job_search_section(self, parent):
        """Create the job search configuration section"""
        search_frame = ttk.LabelFrame(parent, text="🔍 Job Search Configuration", padding="15")
        search_frame.grid(row=2, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Job keywords
        ttk.Label(search_frame, text="Job Keywords:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.keywords_var = tk.StringVar(value="python developer")
        keywords_entry = ttk.Entry(search_frame, textvariable=self.keywords_var, width=40)
        keywords_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(10, 0), pady=5)
        
        # Location
        ttk.Label(search_frame, text="Location:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.location_var = tk.StringVar(value="remote")
        location_entry = ttk.Entry(search_frame, textvariable=self.location_var, width=40)
        location_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(10, 0), pady=5)
        
        # Job type
        ttk.Label(search_frame, text="Job Type:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.job_type_var = tk.StringVar(value="Full-time")
        job_type_combo = ttk.Combobox(search_frame, textvariable=self.job_type_var, 
                                     values=["Full-time", "Part-time", "Contract", "Internship", "Remote"])
        job_type_combo.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(10, 0), pady=5)
        
        # Experience level
        ttk.Label(search_frame, text="Experience Level:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.experience_var = tk.StringVar(value="Entry level")
        experience_combo = ttk.Combobox(search_frame, textvariable=self.experience_var,
                                      values=["Entry level", "Associate", "Mid-Senior level", "Executive"])
        experience_combo.grid(row=3, column=1, sticky=(tk.W, tk.E), padx=(10, 0), pady=5)
        
        search_frame.columnconfigure(1, weight=1)
        
    def create_automation_controls(self, parent):
        """Create the automation control buttons"""
        controls_frame = ttk.LabelFrame(parent, text="🎮 Automation Controls", padding="15")
        controls_frame.grid(row=3, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Main automation button
        self.start_button = ttk.Button(controls_frame, 
                                     text="🚀 Start LinkedIn Automation", 
                                     command=self.start_automation)
        self.start_button.pack(pady=10)
        
        # Secondary controls
        secondary_frame = ttk.Frame(controls_frame)
        secondary_frame.pack(pady=(10, 0))
        
        ttk.Button(secondary_frame, text="⏸️ Pause", command=self.pause_automation).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(secondary_frame, text="⏹️ Stop", command=self.stop_automation).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(secondary_frame, text="🔄 Reset", command=self.reset_automation).pack(side=tk.LEFT)
        
    def create_status_section(self, parent):
        """Create the status and progress section"""
        status_frame = ttk.LabelFrame(parent, text="📊 Status & Progress", padding="15")
        status_frame.grid(row=4, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Status label
        self.status_var = tk.StringVar(value="Ready to start automation")
        status_label = ttk.Label(status_frame, textvariable=self.status_var, font=("Arial", 10, "bold"))
        status_label.pack(pady=(0, 10))
        
        # Progress bar
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate', length=400)
        self.progress.pack(pady=(0, 10))
        
        # Current step
        self.current_step_var = tk.StringVar(value="Waiting to start...")
        step_label = ttk.Label(status_frame, textvariable=self.current_step_var, font=("Arial", 9))
        step_label.pack()
        
        status_frame.columnconfigure(0, weight=1)
        
    def create_results_section(self, parent):
        """Create the results display section"""
        results_frame = ttk.LabelFrame(parent, text="📋 Results & Jobs Found", padding="15")
        results_frame.grid(row=5, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Results text
        self.results_text = scrolledtext.ScrolledText(results_frame, height=8, width=80, wrap=tk.WORD)
        self.results_text.pack(fill=tk.BOTH, expand=True)
        
        # Results actions
        results_actions_frame = ttk.Frame(results_frame)
        results_actions_frame.pack(pady=(10, 0))
        
        ttk.Button(results_actions_frame, text="📥 Export Results", command=self.export_results).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(results_actions_frame, text="🗑️ Clear Results", command=self.clear_results).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(results_actions_frame, text="📊 View Statistics", command=self.show_statistics).pack(side=tk.LEFT)
        
        results_frame.columnconfigure(0, weight=1)
        
    def create_settings_section(self, parent):
        """Create the settings section"""
        settings_frame = ttk.LabelFrame(parent, text="⚙️ Settings", padding="15")
        settings_frame.grid(row=6, column=0, columnspan=2, pady=(0, 15), sticky=(tk.W, tk.E))
        
        # Settings grid
        ttk.Label(settings_frame, text="Auto-apply to jobs:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.auto_apply_var = tk.BooleanVar(value=False)
        auto_apply_check = ttk.Checkbutton(settings_frame, text="Enable automatic job applications", 
                                          variable=self.auto_apply_var)
        auto_apply_check.grid(row=0, column=1, sticky=tk.W, padx=(10, 0), pady=5)
        
        ttk.Label(settings_frame, text="Max jobs to process:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.max_jobs_var = tk.StringVar(value="10")
        max_jobs_entry = ttk.Entry(settings_frame, textvariable=self.max_jobs_var, width=10)
        max_jobs_entry.grid(row=1, column=1, sticky=tk.W, padx=(10, 0), pady=5)
        
        ttk.Label(settings_frame, text="Delay between actions (seconds):").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.delay_var = tk.StringVar(value="2")
        delay_entry = ttk.Entry(settings_frame, textvariable=self.delay_var, width=10)
        delay_entry.grid(row=2, column=1, sticky=tk.W, padx=(10, 0), pady=5)
        
        # Save settings button
        ttk.Button(settings_frame, text="💾 Save Settings", command=self.save_settings).grid(row=3, column=0, columnspan=2, pady=(15, 0))
        
        settings_frame.columnconfigure(1, weight=1)
        
    def browse_resume(self):
        """Browse and select resume file"""
        file_path = filedialog.askopenfilename(
            title="Select Resume File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.resume_path = file_path
            self.resume_path_var.set(file_path)
            self.load_resume_content()
            
    def load_resume_content(self):
        """Load and display resume content"""
        if not self.resume_path:
            return
            
        try:
            doc = Document(self.resume_path)
            self.resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            self.resume_preview.delete(1.0, tk.END)
            self.resume_preview.insert(1.0, self.resume_text)
            
            # Update status
            self.status_var.set(f"Resume loaded: {os.path.basename(self.resume_path)}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load resume: {str(e)}")
            
    def edit_resume(self):
        """Open resume for editing"""
        if not self.resume_path:
            messagebox.showwarning("Warning", "Please select a resume file first")
            return
            
        # Create a simple editor window
        self.create_resume_editor()
        
    def create_resume_editor(self):
        """Create a resume editing window"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Edit Resume")
        editor_window.geometry("600x500")
        
        # Resume content editor
        ttk.Label(editor_window, text="Edit your resume content:").pack(pady=10)
        
        editor_text = scrolledtext.ScrolledText(editor_window, height=20, width=70, wrap=tk.WORD)
        editor_text.pack(padx=20, pady=10, fill=tk.BOTH, expand=True)
        editor_text.insert(1.0, self.resume_text)
        
        # Buttons
        button_frame = ttk.Frame(editor_window)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text="💾 Save Changes", 
                  command=lambda: self.save_resume_changes(editor_text, editor_window)).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="❌ Cancel", 
                  command=editor_window.destroy).pack(side=tk.LEFT)
        
    def save_resume_changes(self, editor_text, editor_window):
        """Save changes made in the resume editor"""
        try:
            new_content = editor_text.get(1.0, tk.END)
            
            # Create new document
            doc = Document()
            for line in new_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Save to file
            doc.save(self.resume_path)
            
            # Update preview
            self.resume_text = new_content
            self.refresh_resume_preview()
            
            messagebox.showinfo("Success", "Resume updated successfully!")
            editor_window.destroy()
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save resume: {str(e)}")
            
    def save_resume(self):
        """Save resume to a new location"""
        if not self.resume_path:
            messagebox.showwarning("Warning", "No resume to save")
            return
            
        save_path = filedialog.asksaveasfilename(
            title="Save Resume As",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if save_path:
            try:
                shutil.copy2(self.resume_path, save_path)
                messagebox.showinfo("Success", f"Resume saved to: {save_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save resume: {str(e)}")
                
    def refresh_resume_preview(self):
        """Refresh the resume preview"""
        if self.resume_path:
            self.load_resume_content()
        else:
            self.resume_preview.delete(1.0, tk.END)
            self.resume_preview.insert(1.0, "No resume selected")
            
    def start_automation(self):
        """Start the LinkedIn automation process"""
        if self.is_automation_running:
            messagebox.showwarning("Warning", "Automation is already running!")
            return
            
        if not self.resume_path:
            messagebox.showwarning("Warning", "Please select a resume file first!")
            return
            
        # Validate inputs
        if not self.keywords_var.get().strip():
            messagebox.showwarning("Warning", "Please enter job keywords!")
            return
            
        # Start automation in separate thread
        self.is_automation_running = True
        self.start_button.config(state='disabled')
        self.progress.start()
        
        thread = threading.Thread(target=self._run_automation)
        thread.daemon = True
        thread.start()
        
    def _run_automation(self):
        """Run the automation process (called in separate thread)"""
        try:
            keywords = self.keywords_var.get()
            location = self.location_var.get()
            
            self.root.after(0, lambda: self.status_var.set(f"Starting LinkedIn automation..."))
            self.root.after(0, lambda: self.current_step_var.set("Initializing browser... If a puzzle appears, complete it in the browser; automation will resume automatically."))
            
            # Start the automation
            success = self.puppeteer_bridge.start_linkedin_automation(keywords, location, self.resume_path)
            
            if success:
                self.root.after(0, lambda: self.status_var.set("Automation completed successfully!"))
                self.root.after(0, lambda: self.current_step_var.set("Loading results..."))
                
                # Get jobs from file
                jobs = self.puppeteer_bridge.get_jobs_from_file()
                
                if jobs:
                    self.root.after(0, lambda: self._display_results(jobs))
                    self.root.after(0, lambda: self.status_var.set(f"Found {len(jobs)} jobs!"))
                else:
                    self.root.after(0, lambda: self.status_var.set("No jobs found"))
                    self.root.after(0, lambda: self.results_text.insert(tk.END, 
                        "No jobs were found during this automation run.\n\nThis could be due to:\n" +
                        "- No matching jobs for the search criteria\n" +
                        "- LinkedIn's page structure changed\n" +
                        "- Automation was interrupted\n\n" +
                        "Try running again or check the browser window for any issues."))
            else:
                self.root.after(0, lambda: self.status_var.set("Automation failed. Check console for details."))
                self.root.after(0, lambda: self.results_text.insert(tk.END, 
                    "Automation failed. Check the console output above for error details.\n\n" +
                    "Common issues:\n- LinkedIn login failed\n- Page elements not found\n" +
                    "- Network connectivity issues\n\nTry running again or check your credentials."))
                
        except Exception as e:
            self.root.after(0, lambda: self.status_var.set(f"Error: {str(e)}"))
            self.root.after(0, lambda: self.results_text.insert(tk.END, 
                f"An error occurred: {str(e)}\n\nPlease check the console for more details."))
            
        finally:
            # Re-enable start button and stop progress
            self.root.after(0, lambda: self.start_button.config(state='normal'))
            self.root.after(0, lambda: self.progress.stop())
            self.root.after(0, lambda: self.current_step_var.set("Ready"))
            self.is_automation_running = False
            
    def pause_automation(self):
        """Pause the automation"""
        if self.is_automation_running:
            self.status_var.set("Automation paused")
            self.current_step_var.set("Paused by user")
            
    def stop_automation(self):
        """Stop the automation"""
        if self.is_automation_running:
            self.puppeteer_bridge.stop_automation()
            self.is_automation_running = False
            self.status_var.set("Automation stopped")
            self.current_step_var.set("Stopped by user")
            self.start_button.config(state='normal')
            self.progress.stop()
            
    def reset_automation(self):
        """Reset the automation state"""
        self.stop_automation()
        self.status_var.set("Ready to start automation")
        self.current_step_var.set("Waiting to start...")
        self.results_text.delete(1.0, tk.END)
        
    def _display_results(self, jobs):
        """Display the job results in the text area"""
        self.results_text.delete(1.0, tk.END)
        
        if not jobs:
            self.results_text.insert(tk.END, "No jobs found.")
            return
            
        self.results_text.insert(tk.END, f"Found {len(jobs)} jobs:\n\n")
        
        for i, job in enumerate(jobs, 1):
            self.results_text.insert(tk.END, f"Job {i}:\n")
            self.results_text.insert(tk.END, f"Title: {job.get('title', 'N/A')}\n")
            self.results_text.insert(tk.END, f"Company: {job.get('company', 'N/A')}\n")
            self.results_text.insert(tk.END, f"Location: {job.get('location', 'N/A')}\n")
            self.results_text.insert(tk.END, f"URL: {job.get('url', 'N/A')}\n")
            
            # Show first 200 characters of description
            description = job.get('description', 'N/A')
            if len(description) > 200:
                description = description[:200] + "..."
            self.results_text.insert(tk.END, f"Description: {description}\n")
            
            self.results_text.insert(tk.END, "-" * 50 + "\n\n")
            
    def export_results(self):
        """Export results to a file"""
        if not self.results_text.get(1.0, tk.END).strip():
            messagebox.showwarning("Warning", "No results to export")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="Export Results",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.results_text.get(1.0, tk.END))
                messagebox.showinfo("Success", f"Results exported to: {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export results: {str(e)}")
                
    def clear_results(self):
        """Clear the results display"""
        self.results_text.delete(1.0, tk.END)
        
    def show_statistics(self):
        """Show automation statistics"""
        # This would show stats like jobs found, success rate, etc.
        messagebox.showinfo("Statistics", "Statistics feature coming soon!")
        
    def save_settings(self):
        """Save current settings"""
        settings = {
            'keywords': self.keywords_var.get(),
            'location': self.location_var.get(),
            'job_type': self.job_type_var.get(),
            'experience': self.experience_var.get(),
            'auto_apply': self.auto_apply_var.get(),
            'max_jobs': self.max_jobs_var.get(),
            'delay': self.delay_var.get()
        }
        
        try:
            with open('gui_settings.json', 'w') as f:
                json.dump(settings, f, indent=2)
            messagebox.showinfo("Success", "Settings saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save settings: {str(e)}")
            
    def load_saved_settings(self):
        """Load previously saved settings"""
        try:
            if os.path.exists('gui_settings.json'):
                with open('gui_settings.json', 'r') as f:
                    settings = json.load(f)
                    
                # Apply saved settings
                self.keywords_var.set(settings.get('keywords', 'python developer'))
                self.location_var.set(settings.get('location', 'remote'))
                self.job_type_var.set(settings.get('job_type', 'Full-time'))
                self.experience_var.set(settings.get('experience', 'Entry level'))
                self.auto_apply_var.set(settings.get('auto_apply', False))
                self.max_jobs_var.set(settings.get('max_jobs', '10'))
                self.delay_var.set(settings.get('delay', '2'))
                
        except Exception as e:
            print(f"Failed to load settings: {e}")

def main():
    """Main function to run the GUI"""
    root = tk.Tk()
    
    # Apply some styling
    style = ttk.Style()
    style.theme_use('clam')
    
    app = LinkedInJobApplierGUI(root)
    
    # Center the window
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    # Make window resizable
    root.resizable(True, True)
    
    root.mainloop()

if __name__ == "__main__":
    main()
